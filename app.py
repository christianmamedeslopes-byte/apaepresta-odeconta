"""
Prestação de Contas – APAE Sidrolândia
Streamlit app para preenchimento facilitado do documento Word.
O arquivo template (template/prestacao_contas.docx) NÃO é alterado.
"""

import io
import copy
import streamlit as st
from docx import Document
import pandas as pd

# ──────────────────────────────────────────────────────────────────────────────
# Helpers de substituição de texto (preservam formatação)
# ──────────────────────────────────────────────────────────────────────────────

def _replace_in_paragraph(para, old: str, new: str) -> bool:
    """Substitui `old` por `new` num parágrafo, tentando preservar a formatação."""
    if old not in para.text:
        return False
    # Caso simples: texto inteiro dentro de um único run
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Caso complexo: texto espalhado por múltiplos runs
    # Concatena tudo no primeiro run e limpa os demais
    full = para.text
    if old in full:
        new_full = full.replace(old, new, 1)
        if para.runs:
            para.runs[0].text = new_full
            for run in para.runs[1:]:
                run.text = ""
        return True
    return False


def replace_in_cell(cell, old: str, new: str):
    """Substitui texto em todos os parágrafos de uma célula."""
    for para in cell.paragraphs:
        _replace_in_paragraph(para, old, new)


def replace_globally(doc, old: str, new: str):
    """Substitui em toda a tabela e parágrafos do documento."""
    if not old or old == new:
        return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, old, new)
    for para in doc.paragraphs:
        _replace_in_paragraph(para, old, new)


def set_cell_text(cell, new_text: str):
    """
    Reescreve o conteúdo de uma célula inteira, preservando o primeiro run.
    Usado para células que contêm APENAS o valor (sem rótulo).
    """
    for i, para in enumerate(cell.paragraphs):
        if para.runs:
            para.runs[0].text = new_text if i == 0 else ""
            for run in para.runs[1:]:
                run.text = ""
        elif i == 0:
            para.add_run(new_text)


# ──────────────────────────────────────────────────────────────────────────────
# Configuração do Streamlit
# ──────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Prestação de Contas – APAE Sidrolândia",
    page_icon="📋",
    layout="wide",
)

st.title("📋 Prestação de Contas – APAE Sidrolândia")
st.caption(
    "Preencha os campos variáveis abaixo e clique em **Gerar Documento**. "
    "O template Word original **não é alterado**."
)

# ──────────────────────────────────────────────────────────────────────────────
# Sidebar – dados fixos (apenas leitura)
# ──────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.header("🔒 Dados Fixos")
    st.markdown("**Entidade:** APAE – Associação de Pais e Amigos dos Excepcionais de Sidrolândia MS")
    st.markdown("**CNPJ:** 33.153.156/0001-61")
    st.markdown("**UF:** MS")
    st.divider()
    st.markdown("**Contador:** Tulio Augusto Gimelli")
    st.markdown("CRC/MS-4051")
    st.divider()
    st.markdown("**Responsável:** Sandra Ione Straliotto Spohr")
    st.markdown("CPF: 322.032.171-20")
    st.divider()
    st.markdown("**Presidente:** Flavio Rodrigues de Sousa")
    st.markdown("CPF: 761.909.606-00")

# ──────────────────────────────────────────────────────────────────────────────
# Seção 1 – Identificação
# ──────────────────────────────────────────────────────────────────────────────

st.subheader("🔖 Identificação")

col1, col2, col3, col4, col5 = st.columns(5)
with col1:
    data_doc = st.text_input("Data do Documento", value="09/04/2026",
                             help="Formato dd/mm/aaaa")
with col2:
    processo = st.text_input("Nº do Processo", value="004/2026")
with col3:
    parceria = st.text_input("Nº da Parceria/Decreto/Ano", value="111/2025")
with col4:
    parcela = st.text_input("Nº da Parcela", value="02/10")
with col5:
    exercicio = st.text_input("Exercício", value="2026")

st.divider()

# ──────────────────────────────────────────────────────────────────────────────
# Seção 2 – Anexo 1: Execução Física
# ──────────────────────────────────────────────────────────────────────────────

st.subheader("📊 Anexo 1 – Execução Física")

col1, col2, col3, col4 = st.columns(4)
with col1:
    qtd_aprovada = st.text_input("Qtd. Aprovada", value="1200")
with col2:
    qtd_reformulada = st.text_input("Qtd. Reformulada", value="")
with col3:
    qtd_exec_periodo = st.text_input("Executada no Período", value="1.328")
with col4:
    qtd_exec_ate = st.text_input("Executada até o Período", value="1.328")

st.divider()

# ──────────────────────────────────────────────────────────────────────────────
# Seção 3 – Anexo 2: Financeiro
# ──────────────────────────────────────────────────────────────────────────────

st.subheader("💰 Anexo 2 – Prestação Financeira")

col1, col2, col3, col4 = st.columns(4)
with col1:
    valor_recebido = st.text_input("Valor Recebido (R$)", value="51.029,57")
with col2:
    rendimentos = st.text_input("Rendimentos Aplicação Financeira (R$)", value="1.338,14")
with col3:
    contrapartida = st.text_input("Valor Contrapartida Utilizado (R$)", value="24.500,00")
with col4:
    total_geral = st.text_input("Total 8.3 (R$)", value="76.867,71")

st.markdown("**Receita / Despesa / Saldo – por período**")
col1, col2, col3, col4, col5, col6 = st.columns(6)
with col1:
    receita_periodo = st.text_input("Receita no Período", value="76.867,71")
with col2:
    receita_ate = st.text_input("Receita até o Período", value="76.867,71")
with col3:
    despesa_periodo = st.text_input("Despesa no Período", value="48.727,33")
with col4:
    despesa_ate = st.text_input("Despesa até o Período", value="48.727,33")
with col5:
    saldo_periodo = st.text_input("Saldo no Período", value="28.140,38")
with col6:
    saldo_ate = st.text_input("Saldo até o Período", value="28.140,38")

st.markdown("**Totais (linha 13)**")
col1, col2, col3, col4, col5, col6 = st.columns(6)
with col1:
    total_rec_periodo = st.text_input("Total Receita Período", value="76.867,71")
with col2:
    total_rec_ate = st.text_input("Total Receita até Período", value="76.867,71")
with col3:
    total_desp_periodo = st.text_input("Total Despesa Período", value="48.727,33")
with col4:
    total_desp_ate = st.text_input("Total Despesa até Período", value="48.727,33")
with col5:
    total_saldo_periodo = st.text_input("Total Saldo Período", value="28.140,38")
with col6:
    total_saldo_ate = st.text_input("Total Saldo até Período", value="28.140,38")

st.divider()

# ──────────────────────────────────────────────────────────────────────────────
# Seção 4 – Anexo 3: Relação de Pagamentos
# ──────────────────────────────────────────────────────────────────────────────

st.subheader("📑 Anexo 3 – Relação de Pagamentos")

PAGAMENTOS_DEFAULT = [
    {"Nº": "01", "Nome / CPF / Endereço": "ELIANE DOS SANTOS CARVALHO R Minas Gerais, 2295 Sidrolândia MS CPF 123.344.281-91",
     "Tipo": "Holerite", "Número": "171", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.675,42"},
    {"Nº": "02", "Nome / CPF / Endereço": "BIANCA BATISTA MARCUZZO R Paraíba, 981 Sidrolandia MS CPF 018.765.341-09",
     "Tipo": "Holerite", "Número": "172", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.675,42"},
    {"Nº": "03", "Nome / CPF / Endereço": "DOUGLAS GOMES MORIKAWA R Napoleão M Siqueira, 787 Sidrolândia MS CPF 016.789.621-04",
     "Tipo": "Holerite", "Número": "76", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "4.460,09"},
    {"Nº": "04", "Nome / CPF / Endereço": "MARILUCIA M SCHNEIDER R Travessa um, 166 casa 05 Sidrolândia MS CPF 576.741.759-87",
     "Tipo": "Holerite", "Número": "108", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "2.299,32"},
    {"Nº": "05", "Nome / CPF / Endereço": "EUNICE RODRIGUES GARBELOTI R Puriu, 20, Campo Grande MS CPF: 175.449.551-68",
     "Tipo": "Holerite", "Número": "57", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.231,02"},
    {"Nº": "06", "Nome / CPF / Endereço": "PATRICIA FINGLER DA COSTA R shoychi Arakaki casa 18 Sidrolândia MS CPF 010.942.12130",
     "Tipo": "Holerite", "Número": "109", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.669,92"},
    {"Nº": "07", "Nome / CPF / Endereço": "RAFAEL BARBOSA DOS SANTOS R Nioaque, 583 Sidrolândia MS CPF 059.417.791-00",
     "Tipo": "Holerite", "Número": "127", "Data Doc": "02/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.401,74"},
    {"Nº": "08", "Nome / CPF / Endereço": "ELIANE DOS SANTOS CARVALHO R Minas Gerais, 2295 Sidrolândia MS CPF 123.344.281-91",
     "Tipo": "Holerite", "Número": "171", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "2.047,25"},
    {"Nº": "09", "Nome / CPF / Endereço": "BIANCA BATISTA MARCUZZO R Paraíba, 981 Sidrolandia MS CPF 018.765.341-09",
     "Tipo": "Holerite", "Número": "172", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "2.047,25"},
    {"Nº": "10", "Nome / CPF / Endereço": "DOUGLAS GOMES MORIKAWA R Napoleão M Siqueira, 787 Sidrolândia MS CPF 016.789.621-04",
     "Tipo": "Holerite", "Número": "76", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "4.460,09"},
    {"Nº": "11", "Nome / CPF / Endereço": "MARILUCIA M SCHNEIDER R Travessa um, 166 casa 05 Sidrolândia MS CPF 576.741.759-87",
     "Tipo": "Holerite", "Número": "108", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "2.299,32"},
    {"Nº": "12", "Nome / CPF / Endereço": "EUNICE RODRIGUES GARBELOTI R Puriu, 20, Campo Grande MS CPF: 175.449.551-68",
     "Tipo": "Holerite", "Número": "57", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.231,02"},
    {"Nº": "13", "Nome / CPF / Endereço": "PATRICIA FINGLER DA COSTA R shoychi Arakaki casa 18 Sidrolândia MS CPF 010.942.12130",
     "Tipo": "Holerite", "Número": "109", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.669,92"},
    {"Nº": "14", "Nome / CPF / Endereço": "RAFAEL BARBOSA DOS SANTOS R Nioaque, 583 Sidrolândia MS CPF 059.417.791-00",
     "Tipo": "Holerite", "Número": "127", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "3.401,74"},
    {"Nº": "15", "Nome / CPF / Endereço": "EDUARDA FERREIRA SIDES R Paraná, 1993 Centro Sidrolândia MS CPF 059.476.571-47",
     "Tipo": "Holerite", "Número": "189", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "2751,42"},
    {"Nº": "16", "Nome / CPF / Endereço": "MARESSA DINIZ COSTA R Distrito Federal, 61 Centro Sidrolândia MS CPF 077.640.431-84",
     "Tipo": "Holerite", "Número": "190", "Data Doc": "03/2026", "Data Pag": "01/05/2026", "Nat": "C", "Valor R$": "1.540,99"},
]

df_pag = pd.DataFrame(PAGAMENTOS_DEFAULT)
df_editado = st.data_editor(
    df_pag,
    use_container_width=True,
    num_rows="dynamic",
    column_config={
        "Nº": st.column_config.TextColumn("Nº", width=40),
        "Nome / CPF / Endereço": st.column_config.TextColumn("Nome / CPF / Endereço", width=350),
        "Tipo": st.column_config.TextColumn("Tipo", width=80),
        "Número": st.column_config.TextColumn("Número", width=70),
        "Data Doc": st.column_config.TextColumn("Data Doc", width=80),
        "Data Pag": st.column_config.TextColumn("Data Pag", width=90),
        "Nat": st.column_config.TextColumn("Nat", width=40),
        "Valor R$": st.column_config.TextColumn("Valor R$", width=90),
    },
    key="tabela_pagamentos",
)

col1, col2 = st.columns(2)
with col1:
    total_pagamentos = st.text_input("Total Pagamentos (R$)", value="48.727,33")
with col2:
    total_acumulado_pag = st.text_input("Total Acumulado Pagamentos (R$)", value="")

st.divider()

# ──────────────────────────────────────────────────────────────────────────────
# Seção 5 – Conciliação Bancária (Anexo 5)
# ──────────────────────────────────────────────────────────────────────────────

st.subheader("🏦 Anexo 5 – Conciliação Bancária")

col1, col2 = st.columns(2)
with col1:
    saldo_banc_data = st.text_input("Data – Saldo Bancário Investimento", value="01/03/2026")
    saldo_banc_valor = st.text_input("Valor – Saldo Bancário (R$)", value="1.338,14")
with col2:
    saldo_ctrl_data = st.text_input("Data – Saldo Controle Convenente", value="01/03/2026")
    saldo_ctrl_valor = st.text_input("Valor – Saldo Controle (R$)", value="1.338,14")

st.markdown("**Transferências recebidas**")
col1, col2, col3, col4 = st.columns(4)
with col1:
    transf1_data = st.text_input("Data Transferência 1", value="11/03/2026")
with col2:
    transf1_valor = st.text_input("Valor Transferência 1 (R$)", value="24.500,00")
with col3:
    transf2_data = st.text_input("Data Transferência 2", value="26/03/2026")
with col4:
    transf2_valor = st.text_input("Valor Transferência 2 (R$)", value="51.029,57")

col1, col2 = st.columns(2)
with col1:
    saldo_concil_data = st.text_input("Data – Saldo Conciliado", value="07/03/2026")
with col2:
    saldo_concil_valor = st.text_input("Valor – Saldo Conciliado (R$)", value="76.867,71")

st.markdown("**Relação de Transferências**")
col1, col2, col3 = st.columns(3)
with col1:
    rel_data = st.text_input("Data Emissão (Meses)", value="02 E 03/2026")
with col2:
    rel_valor = st.text_input("Valor Transferência (R$)", value="48.727,33")
with col3:
    rel_total = st.text_input("Valor Total (com R$)", value="R$48.727,33")

st.divider()

# ──────────────────────────────────────────────────────────────────────────────
# Geração do Documento
# ──────────────────────────────────────────────────────────────────────────────

if st.button("🖨️ Gerar Documento", type="primary", use_container_width=True):
    try:
        doc = Document("template/prestacao_contas.docx")
        tables = doc.tables

        # ── Tabela 1: Identificação Anexo 1 ────────────────────────────────
        t1 = tables[1]
        replace_in_cell(t1.rows[1].cells[1], "004/2026", processo)
        replace_in_cell(t1.rows[1].cells[2], "111/2025", parceria)
        replace_in_cell(t1.rows[1].cells[3], "02/10", parcela)
        replace_in_cell(t1.rows[1].cells[4], "2026", exercicio)

        # ── Tabela 2: Quantidades Anexo 1 ──────────────────────────────────
        t2 = tables[2]
        replace_in_cell(t2.rows[3].cells[2], "1200", qtd_aprovada)
        if qtd_reformulada:
            set_cell_text(t2.rows[3].cells[3], qtd_reformulada)
        replace_in_cell(t2.rows[3].cells[4], "1.328", qtd_exec_periodo)
        replace_in_cell(t2.rows[3].cells[5], "1.328", qtd_exec_ate)

        # ── Tabela 3: Data Anexo 1 ──────────────────────────────────────────
        replace_in_cell(tables[3].rows[0].cells[0], "09/04/2026", data_doc)

        # ── Tabela 5: Identificação Anexo 2 ────────────────────────────────
        t5 = tables[5]
        replace_in_cell(t5.rows[0].cells[4], "111/2025", parceria)
        replace_in_cell(t5.rows[0].cells[6], "02/10", parcela)
        replace_in_cell(t5.rows[1].cells[0], "004/2026", processo)
        replace_in_cell(t5.rows[1].cells[2], "2026", exercicio)
        replace_in_cell(t5.rows[1].cells[5], "51.029,57", valor_recebido)

        # ── Tabela 6: Rendimentos / Contrapartida / Total ───────────────────
        t6 = tables[6]
        replace_in_cell(t6.rows[0].cells[0], "1.338,14", rendimentos)
        replace_in_cell(t6.rows[0].cells[1], "24.500,00", contrapartida)
        replace_in_cell(t6.rows[0].cells[2], "76.867,71", total_geral)

        # ── Tabela 7: Receita / Despesa / Saldo Anexo 2 ────────────────────
        t7 = tables[7]
        # Linha 2 (dados)
        set_cell_text(t7.rows[2].cells[2], receita_periodo)
        set_cell_text(t7.rows[2].cells[4], receita_ate)
        set_cell_text(t7.rows[2].cells[5], despesa_periodo)
        set_cell_text(t7.rows[2].cells[7], despesa_ate)
        set_cell_text(t7.rows[2].cells[8], saldo_periodo)
        set_cell_text(t7.rows[2].cells[9], saldo_ate)
        # Linha 3 (totais)
        set_cell_text(t7.rows[3].cells[2], total_rec_periodo)
        set_cell_text(t7.rows[3].cells[4], total_rec_ate)
        set_cell_text(t7.rows[3].cells[5], total_desp_periodo)
        set_cell_text(t7.rows[3].cells[7], total_desp_ate)
        set_cell_text(t7.rows[3].cells[8], total_saldo_periodo)
        set_cell_text(t7.rows[3].cells[9], total_saldo_ate)
        # Data
        replace_in_cell(t7.rows[5].cells[0], "09/04/2026", data_doc)

        # ── Tabela 9: Pagamentos Anexo 3 ───────────────────────────────────
        t9 = tables[9]
        # Identificação
        replace_in_cell(t9.rows[0].cells[7], "111/2025", parceria)
        replace_in_cell(t9.rows[2].cells[0], "004/2026", processo)
        replace_in_cell(t9.rows[2].cells[5], "02/10", parcela)
        replace_in_cell(t9.rows[2].cells[9], "2026", exercicio)

        # Rebuild células da tabela de pagamentos (linha 6 = dados)
        pag_row = t9.rows[6]
        nums   = "\n".join(str(r["Nº"]) for _, r in df_editado.iterrows())
        nomes  = "\n".join(str(r["Nome / CPF / Endereço"]) for _, r in df_editado.iterrows())
        tipos  = "\n".join(str(r["Tipo"]) for _, r in df_editado.iterrows())
        numdoc = "\n".join(str(r["Número"]) for _, r in df_editado.iterrows())
        datdoc = "\n".join(str(r["Data Doc"]) for _, r in df_editado.iterrows())
        datpag = "\n".join(str(r["Data Pag"]) for _, r in df_editado.iterrows())
        nats   = "\n".join(str(r["Nat"]) for _, r in df_editado.iterrows())
        vals   = "\n".join(str(r["Valor R$"]) for _, r in df_editado.iterrows())

        set_cell_text(pag_row.cells[0],  nums)
        set_cell_text(pag_row.cells[1],  nomes)
        set_cell_text(pag_row.cells[4],  tipos)
        set_cell_text(pag_row.cells[6],  numdoc)
        set_cell_text(pag_row.cells[7],  datdoc)
        set_cell_text(pag_row.cells[9],  datpag)
        set_cell_text(pag_row.cells[11], nats)
        set_cell_text(pag_row.cells[12], vals)

        # Total e total acumulado
        set_cell_text(t9.rows[7].cells[12], total_pagamentos)
        if total_acumulado_pag:
            set_cell_text(t9.rows[8].cells[12], total_acumulado_pag)

        # Data Anexo 3
        replace_in_cell(tables[10].rows[0].cells[0], "09/04/2026", data_doc)

        # ── Tabela 14: Processo Conciliação ────────────────────────────────
        replace_in_cell(tables[14].rows[0].cells[0], "004/2026", processo)

        # ── Tabela 15: Conciliação Bancária ────────────────────────────────
        t15 = tables[15]
        replace_in_cell(t15.rows[0].cells[8], "004/2026", processo)

        # Saldo bancário
        replace_in_cell(t15.rows[2].cells[3], "01/03/2026", saldo_banc_data)
        replace_in_cell(t15.rows[2].cells[7], "1.338,14", saldo_banc_valor)

        # Saldo controle
        replace_in_cell(t15.rows[3].cells[1], "01/03/2026", saldo_ctrl_data)
        replace_in_cell(t15.rows[3].cells[4], "1.338,14", saldo_ctrl_valor)

        # Transferências (duas datas e dois valores na mesma célula)
        replace_in_cell(t15.rows[4].cells[2], "11/03/2026", transf1_data)
        replace_in_cell(t15.rows[4].cells[2], "26/03/2026", transf2_data)
        replace_in_cell(t15.rows[4].cells[5], "24.500,00", transf1_valor)
        replace_in_cell(t15.rows[4].cells[5], "51.029,57", transf2_valor)

        # Saldo conciliado
        replace_in_cell(t15.rows[5].cells[3], "07/03/2026", saldo_concil_data)
        replace_in_cell(t15.rows[5].cells[7], "76.867,71", saldo_concil_valor)

        # ── Tabela 16: Relação de Transferências ───────────────────────────
        t16 = tables[16]
        replace_in_cell(t16.rows[2].cells[1], "02 E 03/2026", rel_data)
        replace_in_cell(t16.rows[2].cells[4], "48.727,33", rel_valor)
        replace_in_cell(t16.rows[3].cells[4], "R$48.727,33", rel_total)

        # ── Salvar e oferecer download ──────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        nome_arquivo = (
            f"prestacao_contas_parcela{parcela.replace('/', '_')}"
            f"_exercicio{exercicio}.docx"
        )

        st.success("✅ Documento gerado com sucesso! O template original **não foi alterado**.")
        st.download_button(
            label="⬇️ Baixar Documento (.docx)",
            data=buf,
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except FileNotFoundError:
        st.error(
            "❌ Template não encontrado em `template/prestacao_contas.docx`.\n\n"
            "Certifique-se de que o arquivo está no repositório dentro da pasta `template/`."
        )
    except Exception as exc:
        st.error(f"❌ Erro ao gerar o documento: {exc}")
        st.exception(exc)
